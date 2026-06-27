import uuid
import json
import io
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.core.database import get_db
from app.api.deps import get_current_user
from app.models.document import Document
from app.models.transcription import Transcription

router = APIRouter(prefix="/export", tags=["export"])


async def _fetch_doc_with_transcription(doc_id: uuid.UUID, db: AsyncSession):
    result = await db.execute(
        select(Document).where(Document.id == doc_id).options(
            selectinload(Document.missionary),
            selectinload(Document.genre),
            selectinload(Document.images),
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    tr_result = await db.execute(
        select(Transcription).where(
            Transcription.document_id == doc_id,
            Transcription.is_current == True,
        ).order_by(Transcription.created_at)
    )
    transcriptions = tr_result.scalars().all()
    return doc, transcriptions


@router.get("/{doc_id}/txt")
async def export_txt(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    doc, transcriptions = await _fetch_doc_with_transcription(doc_id, db)
    lines = [
        f"Reference: {doc.reference_number}",
        f"Missionary: {doc.missionary.name if doc.missionary else 'Unknown'}",
        f"Year: {doc.year}",
        f"Language: {doc.language}",
        "",
        "=== TRANSCRIPTION ===",
        "",
    ]
    for t in transcriptions:
        text = t.corrected_text or t.raw_text or ""
        lines.append(text)
        lines.append("")

    content = "\n".join(lines)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{doc.reference_number}.txt"'},
    )


@router.get("/{doc_id}/json")
async def export_json(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    doc, transcriptions = await _fetch_doc_with_transcription(doc_id, db)
    data = {
        "reference_number": doc.reference_number,
        "missionary": doc.missionary.name if doc.missionary else None,
        "year": doc.year,
        "month": doc.month,
        "genre": doc.genre.name if doc.genre else None,
        "language": doc.language,
        "title": doc.title,
        "description": doc.description,
        "keywords": doc.keywords,
        "transcriptions": [
            {
                "page": i + 1,
                "raw_text": t.raw_text,
                "corrected_text": t.corrected_text,
                "ocr_engine": t.ocr_engine,
                "confidence": t.ocr_confidence,
                "line_data": t.line_data,
            }
            for i, t in enumerate(transcriptions)
        ],
        "exported_at": datetime.utcnow().isoformat(),
    }
    content = json.dumps(data, indent=2, ensure_ascii=False)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{doc.reference_number}.json"'},
    )


@router.get("/{doc_id}/tei")
async def export_tei(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    doc, transcriptions = await _fetch_doc_with_transcription(doc_id, db)
    missionary_name = doc.missionary.name if doc.missionary else "Unknown"
    genre_name = doc.genre.name if doc.genre else "Document"

    tei_lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<TEI xmlns="http://www.tei-c.org/ns/1.0">',
        "  <teiHeader>",
        "    <fileDesc>",
        f"      <titleStmt><title>{doc.title or doc.reference_number}</title></titleStmt>",
        "      <publicationStmt><p>Missionary Archive</p></publicationStmt>",
        "      <sourceDesc>",
        f"        <msDesc><msIdentifier><idno>{doc.reference_number}</idno></msIdentifier>",
        "          <msContents>",
        f"            <msItem><author>{missionary_name}</author>",
        f"            <docDate when='{doc.year}'>{doc.year}</docDate>",
        f"            <textLang mainLang='{doc.language}'>{doc.language}</textLang>",
        "            </msItem>",
        "          </msContents>",
        "        </msDesc>",
        "      </sourceDesc>",
        "    </fileDesc>",
        "  </teiHeader>",
        "  <text><body>",
    ]

    for i, t in enumerate(transcriptions):
        text = t.corrected_text or t.raw_text or ""
        tei_lines.append(f'    <div n="{i+1}" type="page">')
        for line in text.split("\n"):
            escaped = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            tei_lines.append(f"      <l>{escaped}</l>")
        tei_lines.append("    </div>")

    tei_lines += ["  </body></text>", "</TEI>"]
    content = "\n".join(tei_lines)
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="{doc.reference_number}.tei.xml"'},
    )


@router.get("/{doc_id}/docx")
async def export_docx(doc_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    doc, transcriptions = await _fetch_doc_with_transcription(doc_id, db)
    word = DocxDocument()

    # Title
    title = word.add_heading(doc.title or doc.reference_number, 0)
    title.runs[0].font.color.rgb = RGBColor(0x3B, 0x1F, 0x08)

    # Metadata table
    word.add_heading("Document Information", level=2)
    table = word.add_table(rows=0, cols=2)
    meta = [
        ("Reference", doc.reference_number),
        ("Missionary", doc.missionary.name if doc.missionary else "Unknown"),
        ("Year", str(doc.year)),
        ("Language", doc.language),
        ("Genre", doc.genre.name if doc.genre else "N/A"),
        ("OCR Engine", doc.ocr_engine or "N/A"),
    ]
    for key, val in meta:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = val

    word.add_heading("Transcription", level=2)
    for i, t in enumerate(transcriptions):
        if len(transcriptions) > 1:
            word.add_heading(f"Page {i+1}", level=3)
        text = t.corrected_text or t.raw_text or "(No transcription)"
        word.add_paragraph(text)

    buf = io.BytesIO()
    word.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc.reference_number}.docx"'},
    )
